import streamlit as st
import pymysql
import pandas as pd
import certifi
import io
import datetime
from dateutil.relativedelta import relativedelta
from sqlalchemy import create_engine

# --- SETUP MATPLOTLIB ---
import matplotlib
matplotlib.use('Agg') 
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg
import matplotlib.ticker as ticker

# --- LIBRARY REPORTING ---
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape, portrait
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import cm

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.drawing.image import Image as XLImage

# --- GLOBAL COLORS ---
COLOR_HEADER_BLUE = colors.HexColor("#2F5496")
COLOR_TOTAL_YELLOW = colors.HexColor("#FFD966")
COLOR_ROW_EVEN = colors.HexColor("#F2F2F2")
COLOR_ROW_ODD = colors.white
COLOR_BORDER = colors.HexColor("#000000") 

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Aplikasi Retase & Kubikasi", layout="wide", page_icon="lembu.png")

# --- KONEKSI DATABASE ---
@st.cache_resource(ttl=3600)
def init_engine():
    user = st.secrets["db"]["user"]
    password = st.secrets["db"]["password"]
    host = st.secrets["db"]["host"]
    port = st.secrets["db"]["port"]
    database = st.secrets["db"]["database"]
    db_url = f"mysql+pymysql://{user}:{password}@{host}:{port}/{database}"
    engine = create_engine(db_url, connect_args={"ssl": {"ca": certifi.where()}}, pool_recycle=3600, pool_pre_ping=True)
    return engine

# --- HELPER FUNCTIONS ---
def get_bulan_indonesia(bulan_int):
    return ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"][bulan_int]

def get_hari_indonesia(tanggal):
    try:
        if pd.isnull(tanggal): return "-"
        return {'Monday': 'Senin', 'Tuesday': 'Selasa', 'Wednesday': 'Rabu', 'Thursday': 'Kamis', 'Friday': 'Jum\'at', 'Saturday': 'Sabtu', 'Sunday': 'Minggu'}[tanggal.strftime('%A')]
    except: return "-"

def set_cell_bg_docx(cell, color_hex):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def split_date_range_by_month(start_date, end_date):
    result = []
    current = start_date.replace(day=1)
    while current <= end_date:
        month_start = max(start_date, current)
        next_month = current + relativedelta(months=1)
        month_end = min(end_date, next_month - datetime.timedelta(days=1))
        if month_start <= month_end: result.append((month_start, month_end))
        current = next_month
    return result

# --- CHART GENERATORS ---
def generate_chart_harian(df, width_inch=5.0, height_inch=3.0, fontsize=6):
    if df.empty: return None
    rekap = df.groupby('tanggal')['kubikasi'].sum().reset_index()
    if len(rekap) > 15: width_inch = 7.0
    
    rekap['tanggal_str'] = rekap['tanggal'].apply(lambda x: x.strftime('%d/%m'))
    
    fig = Figure(figsize=(width_inch, height_inch), dpi=150)
    canvas = FigureCanvasAgg(fig)
    ax = fig.add_subplot(111)
    fig.patch.set_facecolor('#333333')
    ax.set_facecolor('#333333')
    
    bars = ax.bar(rekap['tanggal_str'], rekap['kubikasi'], color='#5B9BD5', width=0.6)
    ax.set_title('CHART HARIAN', fontsize=10, fontweight='bold', color='white', pad=10)
    ax.tick_params(axis='x', colors='white', labelsize=fontsize, rotation=45)
    ax.tick_params(axis='y', colors='white', labelsize=fontsize)
    for spine in ax.spines.values(): spine.set_color('#555555')
    ax.yaxis.grid(True, color='#555555', linestyle='-', linewidth=0.5)
    ax.set_axisbelow(True)

    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., h + (h*0.02), f'{h:,.1f}', ha='center', va='bottom', color='white', fontsize=fontsize, fontweight='bold')
    fig.tight_layout()
    buf = io.BytesIO(); fig.savefig(buf, format='png', bbox_inches='tight', facecolor=fig.get_facecolor()); buf.seek(0)
    return buf

def generate_chart_bulanan(df_m, width_inch=6.0, height_inch=3.0):
    if df_m.empty: return None
    fig = Figure(figsize=(width_inch, height_inch), dpi=150)
    canvas = FigureCanvasAgg(fig)
    ax = fig.add_subplot(111)
    fig.patch.set_facecolor('#333333')
    ax.set_facecolor('#333333')
    
    bars = ax.bar(df_m['Bulan'], df_m['Kubikasi'], color='#ED7D31', width=0.4)
    ax.set_title('CHART BULANAN', fontsize=10, fontweight='bold', color='white', pad=10)
    ax.tick_params(axis='x', colors='white', labelsize=8)
    ax.tick_params(axis='y', colors='white', labelsize=8)
    for spine in ax.spines.values(): spine.set_color('#555555')
    ax.yaxis.grid(True, color='#555555', linestyle='-', linewidth=0.5)
    ax.set_axisbelow(True)

    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., h + (h*0.02), f'{h:,.2f}', ha='center', va='bottom', color='white', fontsize=8, fontweight='bold')
    fig.tight_layout()
    buf = io.BytesIO(); fig.savefig(buf, format='png', bbox_inches='tight', facecolor=fig.get_facecolor()); buf.seek(0)
    return buf

# ==========================================
# EXPORT GENERATORS
# ==========================================
def process_data_for_export(conn, lokasi_id, start_date, end_date):
    df = pd.read_sql(f"SELECT * FROM data_retase WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
    if not df.empty:
        df['tanggal'] = pd.to_datetime(df['tanggal'])
        df['hari'] = df['tanggal'].apply(get_hari_indonesia)
        rekap = df.groupby(['tanggal', 'hari', 'isi_per_ret']).agg({'jumlah_retase':'sum', 'kubikasi':'sum'}).reset_index().sort_values('tanggal')
        return df, rekap
    return df, pd.DataFrame()

def get_monthly_data(conn, lokasi_id, start_date_global, end_date_global):
    m_data = []
    curr = start_date_global.replace(day=1)
    end_limit = end_date_global.replace(day=1)
    tot_kb = 0.0
    while curr <= end_limit:
        m = curr.month; y = curr.year
        cursor = conn.cursor(); cursor.execute(f"SELECT SUM(kubikasi) FROM data_retase WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}")
        res = cursor.fetchone()
        kb = float(res[0]) if res and res[0] else 0.0
        if kb > 0:
            m_data.append({'Bulan': f"{get_bulan_indonesia(m)} {y}", 'Kubikasi': kb})
            tot_kb += kb
        curr += relativedelta(months=1)
    return pd.DataFrame(m_data), tot_kb

# --- 1. PDF (DIUPDATE SESUAI BBM STANDARDS) ---
def generate_pdf(conn, lokasi_id, nama_lokasi, start_date_g, end_date_g, judul, material, is_landscape):
    buffer = io.BytesIO()
    
    # Hitung maksimal baris data untuk menentukan tinggi kertas dinamis (khusus 1 Kertas)
    max_rows = 0
    for s_date, e_date in split_date_range_by_month(start_date_g, end_date_g):
        _, rekap_temp = process_data_for_export(conn, lokasi_id, s_date, e_date)
        if len(rekap_temp) > max_rows: max_rows = len(rekap_temp)
        
    if is_landscape:
        # Mode 1 Bulan 1 Kertas: Landscape dengan Tinggi Dinamis menyesuaikan isi tabel
        page_width = landscape(A4)[0]
        needed_height = 350 + (max_rows * 20) # 20px per baris
        page_height = max(landscape(A4)[1], needed_height) 
        page_format = (page_width, page_height)
    else:
        # Mode Standard: Kertas A4 Portrait biasa yang mengalir ke bawah
        page_format = portrait(A4)

    doc = SimpleDocTemplate(buffer, pagesize=page_format, rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20)
    elements = []
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='Title', parent=styles['Heading1'], alignment=TA_CENTER, fontSize=14, fontName='Helvetica-Bold', spaceAfter=2)
    sub_style = ParagraphStyle(name='Sub', parent=styles['Normal'], alignment=TA_CENTER, fontSize=12, fontName='Helvetica-Bold', spaceAfter=2)
    period_style = ParagraphStyle(name='Period', parent=styles['Normal'], alignment=TA_CENTER, fontSize=11, spaceAfter=20)

    for idx, (s_date, e_date) in enumerate(split_date_range_by_month(start_date_g, end_date_g)):
        if idx > 0: elements.append(PageBreak())
        df, rekap = process_data_for_export(conn, lokasi_id, s_date, e_date)
        
        elements.append(Paragraph(f"{judul} {material}", title_style))
        elements.append(Paragraph(f"{nama_lokasi}", sub_style))
        elements.append(Paragraph(f"PERIODE {get_bulan_indonesia(s_date.month).upper()} {s_date.year}", period_style))
        
        t_data = [['NO', 'HARI', 'TANGGAL', 'RETASE', 'ISI/1 RET', 'KUBIKASI']]
        tot_ret = 0; tot_kub = 0.0
        if not rekap.empty:
            for i, row in rekap.iterrows():
                t_data.append([str(i+1), row['hari'], row['tanggal'].strftime('%d/%m/%Y'), f"{int(row['jumlah_retase'])}", f"{row['isi_per_ret']:.1f} m³", f"{row['kubikasi']:.1f} m³"])
                tot_ret += row['jumlah_retase']; tot_kub += row['kubikasi']
        else: t_data.append(['-', '-', '-', '-', '-', '-'])

        t_data.extend([['', '', '', 'RINCIAN', '', ''], ['', '', '', 'TOTAL RETASE', f"{int(tot_ret)} Ret", ''], ['', '', '', 'TOTAL KUBIKASI', f"{tot_kub:,.2f} m³", '']])
        
        col_w = [40, 70, 70, 80, 70, 90] if is_landscape else [30, 60, 70, 90, 70, 100]
        t = Table(t_data, colWidths=col_w)
        
        t_style = [
            ('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER), 
            ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_BLUE), 
            ('TEXTCOLOR', (0,0), (-1,0), colors.white), 
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'), 
            ('ALIGN', (0,0), (-1,-1), 'CENTER'), 
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), 
            ('FONTSIZE', (0,0), (-1,-1), 8 if is_landscape else 9),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]
        for i in range(1, len(t_data)-3): t_style.append(('BACKGROUND', (0, i), (-1, i), COLOR_ROW_EVEN if i % 2 == 0 else COLOR_ROW_ODD))
        t_style.extend([('SPAN', (0, -3), (2, -1)), ('SPAN', (3, -3), (5, -3)), ('BACKGROUND', (3, -3), (5, -3), colors.HexColor("#D9E1F2")), ('FONTNAME', (3, -3), (5, -1), 'Helvetica-Bold'), ('SPAN', (4, -2), (5, -2)), ('SPAN', (4, -1), (5, -1))])
        t.setStyle(TableStyle(t_style))
        
        img_buf = generate_chart_harian(df, width_inch=4.5 if is_landscape else 6.0, height_inch=3.0)
        img_rl = RLImage(img_buf, width=350 if is_landscape else 420, height=200) if img_buf else Paragraph("Tidak ada data", period_style)

        if is_landscape:
            # Mode 1 Kertas: Kertas diperpanjang otomatis, jadi aman Kiri Kanan.
            layout_data = [[t, img_rl]]
            layout_table = Table(layout_data, colWidths=[450, 360], style=[('VALIGN', (0,0), (-1,-1), 'TOP')])
            elements.append(layout_table)
        else:
            # Mode Standard Laporan BBM: Mengalir kebawah secara natural (Atas Chart, Bawah Tabel)
            if img_buf:
                elements.append(img_rl)
                elements.append(Spacer(1, 15))
            elements.append(t)
        
    # HALAMAN REKAP BULANAN
    elements.append(PageBreak()); elements.append(Paragraph(f"{judul} {material}", title_style)); elements.append(Paragraph(f"{nama_lokasi}", sub_style)); elements.append(Paragraph("PERBULAN", period_style))
    df_m, tot_all_kb = get_monthly_data(conn, lokasi_id, start_date_g, end_date_g)
    
    # Sesuai permintaan: Atas Grafik, Bawah Tabel Datanya
    img_m_buf = generate_chart_bulanan(df_m, width_inch=6.5 if is_landscape else 5.5, height_inch=3.0)
    if img_m_buf: elements.append(RLImage(img_m_buf, width=450 if is_landscape else 400, height=200)); elements.append(Spacer(1, 15))
        
    tm_data = [['NO', 'BULAN', 'KUBIKASI']]
    for i, r in df_m.iterrows(): tm_data.append([str(i+1), r['Bulan'], f"{r['Kubikasi']:,.2f} m³"])
    tm_data.append(['', 'TOTAL', f"{tot_all_kb:,.2f} m³"])
    t_m = Table(tm_data, colWidths=[50, 200, 200] if is_landscape else [40, 150, 150])
    t_m.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER), ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_BLUE), ('TEXTCOLOR', (0,0), (-1,0), colors.white), ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('FONTNAME', (1,-1), (-1,-1), 'Helvetica-Bold'), ('BACKGROUND', (0,-1), (-1,-1), COLOR_TOTAL_YELLOW)]))
    elements.append(t_m); doc.build(elements); buffer.seek(0)
    return buffer

# --- 2. EXCEL ---
def generate_excel(conn, lokasi_id, nama_lokasi, start_date_g, end_date_g, judul, material):
    output = io.BytesIO(); wb = Workbook(); wb.remove(wb.active); thin = Border(left=Side('thin'), right=Side('thin'), top=Side('thin'), bottom=Side('thin'))
    for s_date, e_date in split_date_range_by_month(start_date_g, end_date_g):
        ws = wb.create_sheet(f"{get_bulan_indonesia(s_date.month)[:3]} {s_date.year}")
        ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 15; ws.column_dimensions['C'].width = 15; ws.column_dimensions['D'].width = 15; ws.column_dimensions['E'].width = 15; ws.column_dimensions['F'].width = 20
        df, rekap = process_data_for_export(conn, lokasi_id, s_date, e_date)
        
        # Center title across Columns A to L to cover both table and chart
        ws.merge_cells('A1:L1'); ws['A1'] = f"{judul} {material}"; ws['A1'].font = Font(bold=True, size=14); ws['A1'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A2:L2'); ws['A2'] = nama_lokasi; ws['A2'].font = Font(bold=True, size=12); ws['A2'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A3:L3'); ws['A3'] = f"PERIODE {get_bulan_indonesia(s_date.month).upper()} {s_date.year}"; ws['A3'].alignment = Alignment(horizontal='center')
        
        r = 5
        headers = ['NO', 'HARI', 'TANGGAL', 'JUMLAH RETASE', 'ISI/1 RET', 'KUBIKASI']
        for i, h in enumerate(headers): c=ws.cell(r, i+1, h); c.border=thin; c.fill=PatternFill("solid", fgColor="2F5496"); c.font=Font(color="FFFFFF", bold=True); c.alignment=Alignment(horizontal='center')
        r += 1
        
        tot_ret = 0; tot_kub = 0.0
        if not rekap.empty:
            for i, row in rekap.iterrows():
                vals = [i+1, row['hari'], row['tanggal'].strftime('%d/%m/%Y'), int(row['jumlah_retase']), f"{row['isi_per_ret']} m³", f"{row['kubikasi']} m³"]
                for j, v in enumerate(vals): c=ws.cell(r, j+1, v); c.border=thin; c.alignment=Alignment(horizontal='center'); c.fill=PatternFill("solid", fgColor="F2F2F2" if i%2!=0 else "FFFFFF")
                tot_ret += row['jumlah_retase']; tot_kub += row['kubikasi']
                r+=1
                
        ws.merge_cells(start_row=r, start_column=4, end_row=r, end_column=6); c=ws.cell(r, 4, "RINCIAN"); c.border=thin; c.fill=PatternFill("solid", fgColor="D9E1F2"); c.font=Font(bold=True); c.alignment=Alignment(horizontal='center'); ws.cell(r,5).border=thin; ws.cell(r,6).border=thin; r+=1
        ws.cell(r,4,"JUMLAH RETASE").border=thin; ws.cell(r,4).font=Font(bold=True); ws.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6); c=ws.cell(r,5,f"{int(tot_ret)} Ret"); c.border=thin; c.font=Font(bold=True); c.alignment=Alignment(horizontal='center'); ws.cell(r,6).border=thin; r+=1
        ws.cell(r,4,"JUMLAH KUBIKASI").border=thin; ws.cell(r,4).font=Font(bold=True); ws.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6); c=ws.cell(r,5,f"{tot_kub:,.2f} m³"); c.border=thin; c.font=Font(bold=True); c.alignment=Alignment(horizontal='center'); ws.cell(r,6).border=thin

        img_buf = generate_chart_harian(df, width_inch=5.0)
        if img_buf: img = XLImage(img_buf); img.width = 450; img.height = 250; ws.add_image(img, 'H5')

    ws2 = wb.create_sheet("Rekap Bulanan"); ws2.column_dimensions['B'].width = 20; ws2.column_dimensions['C'].width = 25
    ws2.merge_cells('A1:C1'); ws2['A1'] = f"{judul} {material}"; ws2['A1'].font = Font(bold=True, size=14); ws2['A1'].alignment = Alignment(horizontal='center')
    ws2.merge_cells('A2:C2'); ws2['A2'] = "PERBULAN"; ws2['A2'].alignment = Alignment(horizontal='center')
    
    df_m, tot_all_kb = get_monthly_data(conn, lokasi_id, start_date_g, end_date_g)
    
    img_m_buf = generate_chart_bulanan(df_m)
    if img_m_buf: img2 = XLImage(img_m_buf); img2.width = 500; img2.height = 250; ws2.add_image(img2, 'A4')
    
    r2 = 18
    for i, h in enumerate(['NO', 'BULAN', 'KUBIKASI']): c=ws2.cell(r2, i+1, h); c.border=thin; c.fill=PatternFill("solid", fgColor="2F5496"); c.font=Font(color="FFFFFF", bold=True); c.alignment=Alignment(horizontal='center')
    r2 += 1
    for i, row in df_m.iterrows():
        vals = [i+1, row['Bulan'], f"{row['Kubikasi']:,.2f} m³"]
        for j, v in enumerate(vals): c=ws2.cell(r2, j+1, v); c.border=thin; c.alignment=Alignment(horizontal='center')
        r2+=1
    ws2.cell(r2, 2, "TOTAL").font=Font(bold=True); ws2.cell(r2, 2).border=thin; ws2.cell(r2, 2).fill=PatternFill("solid", fgColor="FFD966")
    ws2.cell(r2, 3, f"{tot_all_kb:,.2f} m³").font=Font(bold=True); ws2.cell(r2, 3).border=thin; ws2.cell(r2, 3).fill=PatternFill("solid", fgColor="FFD966")
    ws2.cell(r2, 1).border=thin; ws2.cell(r2, 1).fill=PatternFill("solid", fgColor="FFD966")
    
    wb.save(output); output.seek(0)
    return output

# --- 3. WORD ---
def generate_docx(conn, lokasi_id, nama_lokasi, start_date_g, end_date_g, judul, material, is_landscape):
    doc = Document()
    if is_landscape:
        for s in doc.sections:
            s.orientation = 1; s.page_width = Cm(29.7); s.page_height = Cm(21); s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)
    
    for idx, (s_date, e_date) in enumerate(split_date_range_by_month(start_date_g, end_date_g)):
        if idx > 0: doc.add_page_break()
        df, rekap = process_data_for_export(conn, lokasi_id, s_date, e_date)
        
        p = doc.add_paragraph(f"{judul} {material}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.runs[0].font.size=Pt(14); p.paragraph_format.space_after = Pt(2)
        p2 = doc.add_paragraph(f"{nama_lokasi}"); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.runs[0].bold=True; p2.runs[0].font.size=Pt(12); p2.paragraph_format.space_after = Pt(2)
        p3 = doc.add_paragraph(f"PERIODE {get_bulan_indonesia(s_date.month).upper()} {s_date.year}"); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        layout_table = doc.add_table(rows=1, cols=2)
        layout_table.columns[0].width = Cm(14) if is_landscape else Cm(10)
        layout_table.columns[1].width = Cm(12) if is_landscape else Cm(7)
        cell_left = layout_table.cell(0, 0)
        cell_right = layout_table.cell(0, 1)

        t = cell_left.add_table(rows=1, cols=6); t.style = 'Table Grid'
        hdr = t.rows[0].cells
        for i, text in enumerate(['NO', 'HARI', 'TANGGAL', 'JUMLAH RETASE', 'ISI/1 RET', 'KUBIKASI']):
            hdr[i].text = text; set_cell_bg_docx(hdr[i], "2F5496"); hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255); hdr[i].paragraphs[0].runs[0].font.size=Pt(8); hdr[i].paragraphs[0].runs[0].font.bold=True; hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        tot_ret = 0; tot_kub = 0.0
        if not rekap.empty:
            for i, row in rekap.iterrows():
                r = t.add_row().cells
                r[0].text = str(i+1); r[1].text = row['hari']; r[2].text = row['tanggal'].strftime('%d/%m/%Y'); r[3].text = str(int(row['jumlah_retase'])); r[4].text = f"{row['isi_per_ret']} m³"; r[5].text = f"{row['kubikasi']} m³"
                for c in r: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraphs[0].runs[0].font.size=Pt(8)
                tot_ret += row['jumlah_retase']; tot_kub += row['kubikasi']
        
        r1 = t.add_row().cells; r1[3].text = "RINCIAN"; r1[3].merge(r1[5]); set_cell_bg_docx(r1[3], "D9E1F2"); r1[3].paragraphs[0].runs[0].font.bold=True; r1[3].paragraphs[0].runs[0].font.size=Pt(8); r1[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = t.add_row().cells; r2[3].text = "TOTAL RETASE"; r2[4].text = f"{int(tot_ret)} Ret"; r2[4].merge(r2[5]); r2[3].paragraphs[0].runs[0].font.bold=True; r2[4].paragraphs[0].runs[0].font.bold=True; r2[3].paragraphs[0].runs[0].font.size=Pt(8); r2[4].paragraphs[0].runs[0].font.size=Pt(8); r2[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = t.add_row().cells; r3[3].text = "TOTAL KUBIKASI"; r3[4].text = f"{tot_kub:,.2f} m³"; r3[4].merge(r3[5]); r3[3].paragraphs[0].runs[0].font.bold=True; r3[4].paragraphs[0].runs[0].font.bold=True; r3[3].paragraphs[0].runs[0].font.size=Pt(8); r3[4].paragraphs[0].runs[0].font.size=Pt(8); r3[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        img_buf = generate_chart_harian(df, width_inch=4.5 if is_landscape else 3.5, height_inch=3.0)
        if img_buf: 
            p_img = cell_right.add_paragraph()
            p_img.add_run().add_picture(img_buf, width=Cm(12) if is_landscape else Cm(8))

    doc.add_page_break()
    p = doc.add_paragraph(f"{judul} {material}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.runs[0].font.size=Pt(14); p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph("PERBULAN"); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    df_m, tot_all_kb = get_monthly_data(conn, lokasi_id, start_date_g, end_date_g)
    
    img_m_buf = generate_chart_bulanan(df_m, width_inch=6.5, height_inch=3.0)
    if img_m_buf: doc.add_paragraph().add_run().add_picture(img_m_buf, width=Cm(16)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t_m = doc.add_table(rows=1, cols=3); t_m.style = 'Table Grid'
    hdr = t_m.rows[0].cells
    for i, text in enumerate(['NO', 'BULAN', 'KUBIKASI']):
        hdr[i].text = text; set_cell_bg_docx(hdr[i], "2F5496"); hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255); hdr[i].paragraphs[0].runs[0].font.bold=True; hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in df_m.iterrows():
        r = t_m.add_row().cells; r[0].text = str(i+1); r[1].text = row['Bulan']; r[2].text = f"{row['Kubikasi']:,.2f} m³"
        for c in r: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = t_m.add_row().cells; r_t[1].text = "TOTAL"; r_t[2].text = f"{tot_all_kb:,.2f} m³"; r_t[1].paragraphs[0].runs[0].font.bold=True; r_t[2].paragraphs[0].runs[0].font.bold=True; r_t[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for c in r_t: set_cell_bg_docx(c, "FFD966")
        
    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# --- MAIN APP ---
def main():
    if "edit_id" not in st.session_state: st.session_state.edit_id = None
    if "is_super_admin" not in st.session_state: st.session_state.is_super_admin = False

    try: 
        engine = init_engine(); conn = engine.raw_connection(); cursor = conn.cursor() 
        cursor.execute("""CREATE TABLE IF NOT EXISTS lokasi_proyek (id INT AUTO_INCREMENT PRIMARY KEY, nama_tempat VARCHAR(255), kunci_lokasi VARCHAR(255), judul_laporan VARCHAR(255) DEFAULT 'JUMLAH RETASE DAN KUBIKASI OB', jenis_material VARCHAR(255) DEFAULT 'TANAH')""")
        cursor.execute("""CREATE TABLE IF NOT EXISTS data_retase (id INT AUTO_INCREMENT PRIMARY KEY, lokasi_id INT, tanggal DATE, jumlah_retase INT, isi_per_ret FLOAT, kubikasi FLOAT, keterangan TEXT)""")
        cursor.execute("""CREATE TABLE IF NOT EXISTS log_aktivitas (id INT AUTO_INCREMENT PRIMARY KEY, lokasi_id INT, tanggal DATETIME DEFAULT CURRENT_TIMESTAMP, kategori VARCHAR(50), deskripsi TEXT, affected_ids TEXT)""")
        conn.commit()
    except Exception as e: st.error(f"Database Error: {e}"); st.stop()

    if "active_project_id" not in st.session_state: st.session_state.active_project_id = None

    if st.session_state.active_project_id is None:
        with st.sidebar:
            st.header("Super Admin")
            if not st.session_state.is_super_admin:
                with st.form("admin_login"):
                    admin_pass = st.text_input("Password Admin", type="password")
                    if st.form_submit_button("Masuk Halaman Admin"):
                        if admin_pass == st.secrets.get("admin", {}).get("password", ""): st.session_state.is_super_admin = True; st.rerun()
                        else: st.error("Password Salah!")
            else:
                if st.button("⬅️ Keluar Mode Admin", use_container_width=True): st.session_state.is_super_admin = False; st.rerun()

        if st.session_state.is_super_admin:
            st.title("Halaman Super Admin (Retase)")
            st.warning("⚠️ PERINGATAN: Menghapus lokasi akan memusnahkan SELURUH data transaksi di lokasi tersebut secara PERMANEN dan tidak dapat dikembalikan.")
            try: df_lok = pd.read_sql("SELECT * FROM lokasi_proyek", conn)
            except: df_lok = pd.DataFrame()
            if not df_lok.empty:
                with st.container(border=True):
                    if st.session_state.get("clear_konfirmasi"): st.session_state.input_konfirmasi_hapus = ""; st.session_state.clear_konfirmasi = False
                    lok_to_del = st.selectbox("Pilih Lokasi yang akan dihapus permanen:", df_lok['nama_tempat'])
                    konfirmasi = st.text_input('Ketik "KONFIRMASI":', key="input_konfirmasi_hapus")
                    if st.button("Hapus Lokasi", type="primary"):
                        if konfirmasi == "KONFIRMASI":
                            lok_id_del = df_lok[df_lok['nama_tempat'] == lok_to_del].iloc[0]['id']
                            cursor.execute("DELETE FROM data_retase WHERE lokasi_id=%s", (lok_id_del,)); cursor.execute("DELETE FROM log_aktivitas WHERE lokasi_id=%s", (lok_id_del,)); cursor.execute("DELETE FROM lokasi_proyek WHERE id=%s", (lok_id_del,)); conn.commit()
                            st.success("Berhasil dihapus!"); st.session_state.clear_konfirmasi = True; st.rerun()
                        else: st.error("Teks konfirmasi salah.")
            st.stop()

        st.title("🗂️ Menu Utama - Retase & Kubikasi")
        col_left, col_right = st.columns(2, gap="large")
        with col_left:
            with st.container(border=True):
                st.subheader("📂 Masuk ke Lokasi Proyek")
                try: df_lok = pd.read_sql("SELECT * FROM lokasi_proyek", conn)
                except: df_lok = pd.DataFrame()
                if not df_lok.empty:
                    pilih_nama = st.selectbox("Pilih Lokasi:", df_lok['nama_tempat']); input_pass = st.text_input("Password Lokasi:", type="password")
                    if st.button("Masuk Lokasi", type="primary", use_container_width=True):
                        data_lok = df_lok[df_lok['nama_tempat'] == pilih_nama].iloc[0]
                        if input_pass == data_lok['kunci_lokasi']: st.session_state.active_project_id = int(data_lok['id']); st.session_state.active_project_name = data_lok['nama_tempat']; st.rerun()
                        else: st.error("Password Salah!")
        with col_right:
            with st.container(border=True):
                st.subheader("➕ Buat Lokasi Baru")
                new_name = st.text_input("Nama Lokasi Baru"); new_pass = st.text_input("Buat Password", type="password")
                st.warning("⚠️ Harap ingat dan catat Password Lokasi yang Anda buat. Password tidak dapat diubah setelah lokasi dibuat.")
                if st.button("Simpan Lokasi Baru", use_container_width=True):
                    if new_name and new_pass: cursor.execute("INSERT INTO lokasi_proyek (nama_tempat, kunci_lokasi) VALUES (%s, %s)", (new_name, new_pass)); conn.commit(); st.success("Dibuat!"); st.rerun()
                    else: st.error("Isi semua data!")
        st.stop() 

    # --- DASHBOARD PROJECT ---
    lokasi_id = st.session_state.active_project_id
    nama_proyek = st.session_state.active_project_name
    cursor.execute("SELECT judul_laporan, jenis_material FROM lokasi_proyek WHERE id=%s", (lokasi_id,))
    res_set = cursor.fetchone()
    judul_lap_aktif = res_set[0] if res_set[0] else "JUMLAH RETASE DAN KUBIKASI OB"
    jenis_mat_aktif = res_set[1] if res_set[1] else "TANAH"

    with st.sidebar:
        st.header(f"📍 {nama_proyek}")
        st.caption(f"{judul_lap_aktif} {jenis_mat_aktif}")
        if st.button("⬅️ Kembali ke Menu Utama", use_container_width=True): st.session_state.active_project_id = None; st.rerun()

    df_all = pd.read_sql(f"SELECT * FROM data_retase WHERE lokasi_id={lokasi_id} ORDER BY tanggal DESC", conn)
    
    st.title(f"Dashboard: {nama_proyek}")
    t1, t2, t3 = st.tabs(["📝 Input & History", "📊 Laporan & Grafik", "🖨️ Export Dokumen"])
    
    with t1:
        st.subheader("Input Data Retase")
        with st.form("form_input"):
            c1, c2, c3 = st.columns([2, 1, 1])
            with c1: tg = st.date_input("Tanggal"); kt = st.text_input("Keterangan (Opsional)")
            with c2: jr = st.number_input("Jumlah Retase (Trip)", min_value=1, step=1)
            with c3: ir = st.number_input("Isi / 1 Ret (m³)", min_value=0.1, step=0.1, value=5.0)
            if st.form_submit_button("Simpan Data", type="primary"):
                kubikasi = float(jr) * float(ir)
                cursor.execute("INSERT INTO data_retase (lokasi_id, tanggal, jumlah_retase, isi_per_ret, kubikasi, keterangan) VALUES (%s,%s,%s,%s,%s,%s)", (lokasi_id, tg, jr, ir, kubikasi, kt)); conn.commit(); st.success("Tersimpan!"); st.rerun()
                
        st.divider()
        with st.expander("⏳ RIWAYAT INPUT & EDIT", expanded=True):
            if st.session_state.edit_id:
                st.markdown("### ✏️ Edit Data")
                cursor.execute("SELECT tanggal, jumlah_retase, isi_per_ret, keterangan FROM data_retase WHERE id=%s", (st.session_state.edit_id,))
                res = cursor.fetchone()
                if res:
                    with st.form("edit_form"):
                        c1, c2, c3 = st.columns([2,1,1])
                        with c1: e_tg = st.date_input("Tanggal", value=res[0]); e_kt = st.text_input("Keterangan", value=str(res[3]))
                        with c2: e_jr = st.number_input("Jumlah Retase", value=int(res[1]))
                        with c3: e_ir = st.number_input("Isi / 1 Ret (m³)", value=float(res[2]))
                        ce1, ce2 = st.columns(2)
                        if ce1.form_submit_button("Simpan Perubahan", type="primary"):
                            cursor.execute("UPDATE data_retase SET tanggal=%s, jumlah_retase=%s, isi_per_ret=%s, kubikasi=%s, keterangan=%s WHERE id=%s", (e_tg, e_jr, e_ir, e_jr*e_ir, e_kt, st.session_state.edit_id)); conn.commit(); st.session_state.edit_id = None; st.rerun()
                        if ce2.form_submit_button("Batal"): st.session_state.edit_id = None; st.rerun()
            
            if not df_all.empty:
                c_f1, c_f2 = st.columns(2)
                with c_f1: use_filter = st.checkbox("Gunakan Filter Tanggal Riwayat")
                with c_f2: tgl_filter = st.date_input("Filter Tanggal", disabled=not use_filter)
                
                df_view = df_all.copy()
                if use_filter:
                    df_view = df_view[pd.to_datetime(df_view['tanggal']).dt.date == tgl_filter]
                
                st.write(f"Menampilkan {min(len(df_view), 50)} data terbaru...")
                
                # FITUR CONTAINER AGAR HISTORY TIDAK MEMANJANG KE BAWAH (SCROLLABLE)
                with st.container(height=400, border=True):
                    for _, row in df_view.head(50).iterrows():
                        ca, cb, cc, cd = st.columns([2, 3, 2, 2])
                        ca.write(f"📅 {row['tanggal'].strftime('%d/%m/%Y')}")
                        cb.write(f"**{row['jumlah_retase']} Ret** x {row['isi_per_ret']} m³")
                        cc.write(f"**{row['kubikasi']:,.1f} m³**")
                        with cd:
                            b1, b2 = st.columns(2)
                            if b1.button("✏️", key=f"e_{row['id']}", help="Edit Data"): st.session_state.edit_id = row['id']; st.rerun()
                            if b2.button("❌", key=f"d_{row['id']}", help="Hapus Data"): cursor.execute("DELETE FROM data_retase WHERE id=%s", (row['id'],)); conn.commit(); st.rerun()
                        st.divider()
            else:
                st.info("Belum ada riwayat input.")

    with t2:
        c_p1, c_p2 = st.columns(2)
        if 't2_start' not in st.session_state: st.session_state.t2_start = datetime.date.today().replace(day=1)
        if 't2_end' not in st.session_state: st.session_state.t2_end = datetime.date.today()
        with c_p1: start_rep = st.date_input("Mulai", value=st.session_state.t2_start, key="p1"); st.session_state.t2_start = start_rep
        with c_p2: end_rep = st.date_input("Sampai", value=st.session_state.t2_end, key="p2"); st.session_state.t2_end = end_rep
        
        df_rep = df_all[(pd.to_datetime(df_all['tanggal']).dt.date >= start_rep) & (pd.to_datetime(df_all['tanggal']).dt.date <= end_rep)]
        st.markdown(f"""<div style="background-color:#D9E1F2;padding:15px;border-radius:10px;border:1px solid #2F5496;text-align:center;margin-bottom:20px;"><h2 style="color:#2F5496;margin:0;">TOTAL KUBIKASI: {df_rep['kubikasi'].sum():,.2f} m³</h2><span style="color:#333;font-weight:bold;">Total Retase: {df_rep['jumlah_retase'].sum():,.0f} Trip</span></div>""", unsafe_allow_html=True)

        with st.expander("🛠️ PENGATURAN NAMA LAPORAN & PROYEK", expanded=False):
            with st.form("form_setting"):
                st.caption("Ubah jenis material ke apa saja (contoh: SEMEN, PASIR, TANAH). Teks ini akan langsung tercetak di atas laporan PDF/Excel/Word Anda.")
                s1 = st.text_input("Judul Laporan Utama:", value=judul_lap_aktif)
                s2 = st.text_input("Jenis Material:", value=jenis_mat_aktif)
                s3 = st.text_input("Nama Lokasi / Proyek:", value=nama_proyek)
                if st.form_submit_button("Simpan Pengaturan"): cursor.execute("UPDATE lokasi_proyek SET judul_laporan=%s, jenis_material=%s, nama_tempat=%s WHERE id=%s", (s1, s2, s3, lokasi_id)); conn.commit(); st.session_state.active_project_name = s3; st.rerun()

        st.divider()
        c_g1, c_g2 = st.columns([1, 1])
        with c_g1:
            st.subheader("📊 Grafik Harian")
            st.caption("Sesuai periode tanggal di atas.")
            img_harian = generate_chart_harian(df_rep)
            if img_harian: st.image(img_harian)
            else: st.info("Tidak ada data harian.")
        
        with c_g2:
            st.subheader("📊 Grafik Bulanan")
            st.caption("Sesuai periode tanggal di atas.")
            df_m_tab2, tot_kb_tab2 = get_monthly_data(conn, lokasi_id, start_rep, end_rep)
            img_bulanan = generate_chart_bulanan(df_m_tab2)
            if img_bulanan: st.image(img_bulanan)
            else: st.info("Tidak ada data bulanan.")
            
        if not df_m_tab2.empty:
            st.dataframe(df_m_tab2, use_container_width=True, hide_index=True)

        # TABEL DETAIL DI TAB 2 (Sesuai Request)
        st.divider()
        st.subheader("📋 Detail Data Retase & Kubikasi")
        if not df_rep.empty:
            df_display = df_rep[['tanggal', 'jumlah_retase', 'isi_per_ret', 'kubikasi', 'keterangan']].copy()
            df_display.rename(columns={'tanggal':'Tanggal', 'jumlah_retase':'Jumlah Retase', 'isi_per_ret':'Isi / 1 Ret (m³)', 'kubikasi':'Total Kubikasi (m³)', 'keterangan':'Keterangan'}, inplace=True)
            st.dataframe(df_display, use_container_width=True, hide_index=True)
        else:
            st.info("Tidak ada data pada periode ini.")

    with t3:
        st.header("🖨️ Export Laporan Retase")
        mode_export = st.radio("Mode Tampilan Kertas:", ["Standard", "1 Bulan 1 Kertas"], horizontal=True)
        is_landscape = "1 Kertas" in mode_export
        
        c_d1, c_d2 = st.columns(2)
        if 't3_start' not in st.session_state: st.session_state.t3_start = datetime.date.today().replace(day=1)
        if 't3_end' not in st.session_state: st.session_state.t3_end = datetime.date.today()
        with c_d1: start_date_exp = st.date_input("Dari", value=st.session_state.t3_start, key="e1"); st.session_state.t3_start = start_date_exp
        with c_d2: end_date_exp = st.date_input("Sampai", value=st.session_state.t3_end, key="e2"); st.session_state.t3_end = end_date_exp

        if start_date_exp <= end_date_exp:
            c1, c2, c3 = st.columns(3)
            with c1:
                if st.button("📕 Download PDF", use_container_width=True): 
                    with st.spinner("Membuat PDF..."):
                        pdf = generate_pdf(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, judul_lap_aktif, jenis_mat_aktif, is_landscape)
                        st.download_button("⬇️ Simpan PDF", pdf, f"Retase_{nama_proyek}.pdf", "application/pdf")
            with c2:
                if st.button("📗 Download Excel", use_container_width=True): 
                    with st.spinner("Membuat Excel..."):
                        xl = generate_excel(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, judul_lap_aktif, jenis_mat_aktif)
                        st.download_button("⬇️ Simpan Excel", xl, f"Retase_{nama_proyek}.xlsx")
            with c3:
                if st.button("📘 Download Word", use_container_width=True): 
                    with st.spinner("Membuat Word..."):
                        doc = generate_docx(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, judul_lap_aktif, jenis_mat_aktif, is_landscape)
                        st.download_button("⬇️ Simpan Word", doc, f"Retase_{nama_proyek}.docx")
        else: st.error("Tanggal salah.")

if __name__ == "__main__":
    main()